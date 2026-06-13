import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def map_speakers(results):
    spk_map = {}
    counter = ord('A')
    for item in results:
        orig = item["speaker"]
        if orig not in spk_map:
            spk_map[orig] = chr(counter)
            counter += 1
    return spk_map

def print_and_save_transcript(results, spk_map, output_file="outputs/penlu.txt"):
    os.makedirs("outputs", exist_ok=True)
    print("\n=== 🎯 带说话人笔录 ===")
    with open(output_file, "w", encoding="utf-8") as f:
        for item in results:
            label = spk_map.get(item["speaker"], item["speaker"])
            line = f"[{label}] {item['start']:.1f}s - {item['end']:.1f}s: {item['text']}"
            print(line)
            f.write(line + "\n")
    print(f"\n✅ 笔录已保存: {output_file}")

def export_to_word(results, spk_map, filename=None):
    if filename is None:
        filename = f"outputs/笔录_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    os.makedirs("outputs", exist_ok=True)
    doc = Document()
    
    title = doc.add_heading('公安智能语音笔录', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph('=' * 50)
    
    for item in results:
        label = spk_map.get(item["speaker"], item["speaker"])
        p = doc.add_paragraph()
        p.add_run(f'[{label}] ').bold = True
        p.add_run(f'{item["start"]::.1f}s - {item["end"]::.1f}s: ')
        p.add_run(item["text"])
    
    doc.save(filename)
    print(f"✅ Word 专业笔录已导出: {filename}")
    return filename